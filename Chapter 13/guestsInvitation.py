#! python3

import docx
from docx.shared import Pt

doc = docx.Document()
styleCooperBlack = doc.styles['Normal']
fontCooperBlack = styleCooperBlack.font
fontCooperBlack.name = 'Cooper Black'
fontCooperBlack.size = Pt(15)

styleDate = doc.styles['Normal']
font = styleDate.font
font.name = 'Times New Roman'
font.size = Pt(15)

fileObject = open('D:\\Scripts\\guests.txt','r')
for line in fileObject.readlines():
	paragrapghObject = doc.add_paragraph('It would be a pleasure to have the company of')
	paragrapghObject.style = styleCooperBlack
	
	doc.add_paragraph(line.strip())
	
	doc.add_paragraph('at 11010 Memory Lane on the Evening of')
	
	paragrapghObject1 = doc.add_paragraph('April 1st')
	paragrapghObject1.style = styleDate
	
	doc.add_paragraph('at 7 o\'clock')
	
	doc.add_page_break()
	#doc.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
	
fileObject.close()
doc.save('D:\\Scripts\\invitation1.docx')
